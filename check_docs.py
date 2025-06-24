import os
from PyPDF2 import PdfFileReader
import zipfile
import os
import re
import pdfreader 
from loguru import logger
import docx
import typer

import glob
import subprocess

app = typer.Typer()

#for doc in glob.iglob("*.doc"):
#    subprocess.call(['soffice', '--headless', '--convert-to', 'docx', doc])

def process_cvs_directory(directory: str, verbose: bool = False):
    """Process CVs in the specified directory and extract GitHub/LinkedIn links."""
    # Re-initialize results list for Markdown table
    markdown_table = ["| Filename | GitHub URL | Linkedin | ", "| --- | --- | --- |"]
    
    unzip_dir = directory
    files_processed = 0
    files_with_links = 0
    
    # Loop through unzipped PDF files
    for root, dirs, files in os.walk(unzip_dir):
        for pdf_file in sorted(files):
            pdf_path = os.path.join(root, pdf_file)
            files_processed += 1
            
            if verbose:
                logger.info(f"Processing: {pdf_path}")
            
            if pdf_file.lower().endswith('.doc') or pdf_file.lower().endswith('.docx'):
                text = process_word(pdf_path)
            elif not pdf_file.lower().endswith('.pdf'):
                continue  # Skip non-PDF files
            else:
                text = process_pdf(pdf_path)
            
            # Extract GitHub links, ignoring case
            gitlab = re.findall(r'gitlab\.com/[\w/-]+', text, re.IGNORECASE)
            bitbucket = re.findall(r'bitbucket\.com/[\w/-]+', text, re.IGNORECASE)
            github_links = re.findall(r'github\.com/[\w/-]+', text, re.IGNORECASE)
            linkedin = re.findall(r'linkedin\.com/[\w/-]+', text, re.IGNORECASE)
            if github_links or bitbucket or gitlab:
                files_with_links += 1
                repo = f"{ gitlab[0] if gitlab else '' } {bitbucket[0] if bitbucket else ''} {github_links[0] if github_links else ''}"
                markdown_table.append(f"| {pdf_file} | {repo} | {linkedin[0] if linkedin else ''} | ")

            #if github_links:
            #    print(github_links)
            #    for github_link in github_links:
            #        markdown_table.append(f"| {pdf_file} | {github_link} | | ")
            #if linkedin:
            #    print(linkedin)
            #    for linkedin in github_links:
            #        markdown_table.append(f"| {pdf_file} | | {linkedin} |")

    # Join markdown table rows into a single string
    markdown_table_str = '\n'.join(markdown_table)
    logger.info(f"Processed {files_processed} files, found links in {files_with_links} files")
    print(markdown_table_str)


def process_word(fn):
	# Open the Word document
    try:
        try:
            doc = docx.Document(fn)
        except Exception as ex:
            logger.error(f"{ex}: try converting")
            subprocess.call(['soffice', '--headless', '--convert-to', 'docx', fn])
            doc = docx.Document(fn)

        full_text = []

        # Loop through paragraphs 
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Loop through tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    
        # Join all paragraphs and table texts
        full_text = '\n'.join(full_text)
        return full_text
    except Exception as ex:
        logger.error(ex)
        return ''

def process_pdf(fn):
	# Read PDF content
    try:
        pdf_reader = PdfFileReader(open(fn, 'rb'))
        #pdf_reader = pdfreader.SimplePDFViewer(open(pdf_path, 'rb'))
        pdf_text = ''
        for page_num in range(pdf_reader.getNumPages()):
            page = pdf_reader.getPage(page_num)
            pdf_text += page.extractText()
        return pdf_text
    except:
        return ''


@app.command()
def main(
    directory: str = typer.Argument(default="Applications", help="Directory containing CVs to examine"),
    verbose: bool = typer.Option(False, "-v", "--verbose", help="Show each file being processed")
):
    """Process CVs in the specified directory and extract GitHub/LinkedIn links."""
    logger.info(f"Scanning directory: {directory}")
    process_cvs_directory(directory, verbose)


if __name__ == "__main__":
    app()


